import codecs
import re
from docx import Document
from lxml import etree
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
import sys
import zipfile
import docx.enum.text
import os
from nltk.tokenize import sent_tokenize
import json


def delete_paragraph(paragraph):
    if paragraph is not None:
        p = paragraph._element
        if p is not None:
            p.getparent().remove(p)
            paragraph._p = paragraph._element = None

# 删除docx里面的bookmark
def DeleteAllBookmark(d):
    for element in d.element.body:
        if (element.__class__.__name__ == 'CT_Bookmark') or (element.__class__.__name__ == 'CT_MarkupRange'):
            d.element.body.remove(element)
            print(element.id)
        else:
            for ele in element:
                if (ele.__class__.__name__ == 'CT_Bookmark') or (ele.__class__.__name__ == 'CT_MarkupRange'):
                    element.remove(ele)

"""
对docx处理并获得xml

Parameters:
  param1 - 输入的docx文件名
  param2 - 转换后的xml文件名

Returns:
"""
def processDocx(filename, newfile):
    doc = Document(filename)
    delete_references = False  # 标记删除Acknowledgment之后的段落

    #print("有{}张图片".format(len(doc.inline_shapes)))

    # 删除表格
    #print("有{}个表格".format(len(doc.tables)))
    for table in doc.tables:
        table._element.getparent().remove(table._element)
    DeleteAllBookmark(doc)

    # 处理section
    #print("有{}个section".format(len(doc.sections)))
    # 取消“奇偶页不同”
    doc.settings.odd_and_even_pages_header_footer = False
    # 删除格式
    for section in doc.sections:
        # 取消“首页不同”
        section.different_first_page_header_footer = False
        # 删除页眉、页脚
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')
        if cols:
            cols[0].set(qn('w:num'), '1')

    f = codecs.open(newfile, 'w', 'utf-8')
    for p in doc.paragraphs:
        if len(p.text) <= 2:
            continue
        if p.text.lower() == 'references':  # 删除references之后的内容
            delete_references = True
        if (p.text == "\n" or p.text == '' or p.text.isspace() or delete_references
                or p.text.lower().startswith("table") or p.text.lower().startswith("fig")):
            p.clear()
            delete_paragraph(p)
        # if delete_references or p.text.lower().startswith("table") or p.text.lower().startswith("fig"):
        #     p.clear()
        #     delete_paragraph(p)
        else:
            #print("段落" + p.text)
            p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.LEFT
            p_xml_str = p._p.xml  # 按段落获取xml
            for c in p._element.getiterator():
                tag = c.tag.split('}')[1]
                if tag == 'drawing' or tag == 'inline' or tag == 'textbox':  # 删除带w:drawing的标签（即删除图片）
                    c.getparent().remove(c)
                if tag == 'tab':  # 删除表格
                    #c.getparent().remove(c)
                    delete_paragraph(p)
            f.write(p_xml_str)
    doc.save(filename)

"""
将docx转为txt

Parameters:
  param1 - 输入的docx文件名
  param2 - 该docx对应的配置json

Returns:
"""
def convertToTxt(file, fileJson):
    filename, extension = os.path.splitext(file)
    # 新建和打开txt文档
    f = codecs.open(filename+".txt", 'w', 'utf-8')
    # 打开docx的文档并读入名为file的变量中
    doc = docx.Document(file)
    # 输入docx中的段落数，以检查是否空文档
    #print('段落数:' + str(len(doc.paragraphs)))
    # 将每个段落的内容都写进去txt里面
    for para in doc.paragraphs:
        if len(para.text) <= 2:
            continue
        paragraph = para.text.replace('- ', '').replace('Fig.', 'Fig').replace('Figure.', 'Figure')\
            .replace('Table.', 'Table')
        #print(paragraph)
        f.write(paragraph+'\n')
    f.close()
    processTxt(filename+'.txt')


"""
对获得的txt再进行处理,对换行的句子进行合并

Parameters:
  param1 - 输入的docx文件名
  param2 - 该docx对应的配置json

Returns:
"""
def processTxt(file):
    # 读取txt，合并需要换行的句子
    # contents存储原本txt的句子
    contents = []
    f = codecs.open(file, 'r', 'utf-8')
    lines = f.readlines()
    lines.append(' ')
    if len(lines) > 1:
        i = 1
        while i < len(lines):
            # print(i)
            line = lines[i-1]
            if len(line) > 1 and line[-2] == '.':
                contents.append(line)
            elif len(line) > 1 and line[-2] == '-':
                contents.append(line[:-2] + lines[i])
                i += 1
            else:
                contents.append(line + " " + lines[i])
                i += 1
            i += 1
    f.close()

    # 读取合并后的句子，根据 . ! ? 分句
    # sentences保存处理后的句子
    sentences = sent_tokenize(''.join(contents))
    f = codecs.open(file, 'w', 'utf-8')
    for s in sentences:
        #print(s)
        sentence = s.replace('\n', ' ')
        f.write(sentence + '\n')
    f.close()

# 测试docx的Json
def testJson(jsonObj):
    with codecs.open(r'D:\新建文件夹\testArticle.json', 'r', 'utf8')as fp:
        json_data = json.load(fp)
        # print('这是文件中的json数据：', json_data)
        # print('这是读取到文件数据的数据类型：', type(json_data))

        deleteSentences = json_data.get("DeleteSentences")

        doc = Document(r'D:\新建文件夹\testArticle.docx')
        flag = False
        # sentenceFlag作为处理 json中需要删除句子 的标志位
        sentenceFlag = False
        for p in doc.paragraphs:
            if p.text.startswith(json_data.get("RemainTag")):
                flag = True
            elif p.text.startswith(json_data.get("EliminateTag")):
                flag = False
            for s in deleteSentences:
                if p.text.startswith(s.get("start")) and p.text.endswith(s.get("end")):
                    sentenceFlag = False
                else:
                    sentenceFlag = True
            if flag and sentenceFlag:
                print("段落 " + p.text)

if __name__ == '__main__':
    processDocx(r'D:\新建文件夹\testArticle.docx', r'D:\新建文件夹\testArticle.xml')
    convertToTxt(r'D:\新建文件夹\testArticle.docx', r'D:\新建文件夹\testArticle.json')
